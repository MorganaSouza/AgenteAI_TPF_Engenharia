import google.generativeai as genai
from docx import Document
from openpyxl import Workbook
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
import re
import os
from dotenv import load_dotenv # Importa a função load_dotenv

# --- Configurações da API ---
load_dotenv() # Carrega as variáveis do arquivo .env
API_KEY = os.getenv("GOOGLE_API_KEY")

if not API_KEY:
    raise ValueError("A variável de ambiente GOOGLE_API_KEY não foi definida.")

genai.configure(api_key=API_KEY)

# --- Funções de processamento e salvamento ---

def gerar_resposta(prompt: str) -> str:
    """Interage com o Gemini e retorna a resposta."""
    modelo = genai.GenerativeModel("gemini-1.5-flash")
    resposta = modelo.generate_content(prompt)
    return resposta.text

def extrair_topicos(texto: str) -> list:
    """
    Extrai os tópicos numerados e subtópicos da resposta do Gemini.
    """
    # A expressão regular busca por linhas que começam com um número e um ponto,
    # capturando todo o texto até a próxima linha numerada ou o final.
    topicos_encontrados = re.findall(r'^\s*(\d+\.\s*.*?)(?=\s*\d+\.|\Z)', texto, re.DOTALL | re.MULTILINE)
    
    # Se a extração falhar, retorna o texto original dividido por linhas
    if not topicos_encontrados:
        return [linha.strip() for linha in texto.split("\n") if linha.strip()]
    
    return topicos_encontrados

def salvar_docx(topicos: list, nome="relatorio_ia.docx"):
    """
    Salva os tópicos em um arquivo DOCX com título e parágrafos.
    """
    doc = Document()
    doc.add_heading("Relatório de Inteligência Artificial", 0)

    for topico in topicos:
        doc.add_paragraph(topico)
    
    doc.save(nome)

def salvar_xlsx(topicos: list, nome="relatorio_ia.xlsx"):
    """
    Salva os tópicos em um arquivo XLSX, com cada tópico em uma célula.
    """
    wb = Workbook()
    ws = wb.active
    ws.title = "Relatório IA"

    for i, topico in enumerate(topicos, start=1):
        ws[f"A{i}"] = topico

    wb.save(nome)

def salvar_pdf(topicos: list, nome="relatorio_ia.pdf"):
    """
    Salva os tópicos em um arquivo PDF com título e espaçamento dinâmico.
    """
    # Cria o documento PDF
    doc = SimpleDocTemplate(nome, pagesize=letter)
    
    # Obtém os estilos de parágrafo padrão
    styles = getSampleStyleSheet()
    
    # Cria uma lista para armazenar os elementos do PDF
    flowables = []
    
    # Adiciona o título
    titulo = Paragraph("Relatório de Inteligência Artificial", styles['Heading1'])
    flowables.append(titulo)
    flowables.append(Spacer(1, 12))  # Espaço após o título
    
    # Adiciona os tópicos como parágrafos
    for topico in topicos:
        paragrafo = Paragraph(topico, styles['Normal'])
        flowables.append(paragrafo)
        flowables.append(Spacer(1, 6)) # Espaço entre os tópicos
        
    # Constrói o PDF
    doc.build(flowables)

# --- Execução principal ---

if __name__ == "__main__":
    prompt_estruturado = """
    Você é um assistente especializado em tecnologia. 
    Explique de forma clara e didática o que é Inteligência Artificial.
    A resposta deve ser organizada em tópicos numerados, como:

    1. Definição simples
    2. Principais características
    3. Exemplos práticos
    4. Benefícios
    5. Desafios e limitações
    6. Conclusão

    Não use negrito, itálico ou qualquer outro tipo de formatação além da numeração.
    """
    
    resposta = gerar_resposta(prompt_estruturado)
    topicos_processados = extrair_topicos(resposta)

    print("--- Resposta do Gemini (texto original) ---\n", resposta)
    print("\n--- Tópicos extraídos para formatação ---\n", topicos_processados)

    salvar_docx(topicos_processados)
    salvar_xlsx(topicos_processados)
    salvar_pdf(topicos_processados)

    print("\n--- Arquivos gerados com sucesso! ---")
    print("Verifique os arquivos: relatorio_ia.docx, relatorio_ia.xlsx, relatorio_ia.pdf")